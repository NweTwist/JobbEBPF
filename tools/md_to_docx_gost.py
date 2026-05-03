#!/usr/bin/env python3
"""
Markdown to DOCX converter with GOST-like formatting defaults.

Defaults (GOST 7.32-2017 common practice):
- Font: Times New Roman, 14 pt
- Line spacing: 1.5
- First line indent: 1.25 cm
- Margins: left 3.0 cm, right 1.0 cm, top 2.0 cm, bottom 2.0 cm
- Paragraph alignment: justified
- Headings: bold; level 1 centered, level 2+ left

Usage:
  python tools/md_to_docx_gost.py text.md text.docx

Requires:
  pip install python-docx
"""

from __future__ import annotations

import re
import sys
from pathlib import Path

try:
    from docx import Document
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.shared import Cm, Pt
    from docx.oxml.ns import qn
except Exception as exc:  # pragma: no cover
    print("Missing dependency: python-docx. Install with: pip install python-docx")
    raise SystemExit(2) from exc


DEFAULTS = {
    "font_name": "Times New Roman",
    "font_size_pt": 14,
    "line_spacing": 1.5,
    "first_line_indent_cm": 1.25,
    "margin_left_cm": 3.0,
    "margin_right_cm": 1.0,
    "margin_top_cm": 2.0,
    "margin_bottom_cm": 2.0,
}


HEADING_RE = re.compile(r"^(#{1,6})\s+(.*)$")
TABLE_SEPARATOR_RE = re.compile(r"^\s*\|?\s*[-:]+(?:\s*\|\s*[-:]+)*\s*\|?\s*$")
LIST_RE = re.compile(r"^\s*[-*]\s+(.*)$")


def set_document_style(doc: Document) -> None:
    section = doc.sections[0]
    section.left_margin = Cm(DEFAULTS["margin_left_cm"])
    section.right_margin = Cm(DEFAULTS["margin_right_cm"])
    section.top_margin = Cm(DEFAULTS["margin_top_cm"])
    section.bottom_margin = Cm(DEFAULTS["margin_bottom_cm"])

    style = doc.styles["Normal"]
    style.font.name = DEFAULTS["font_name"]
    style.font.size = Pt(DEFAULTS["font_size_pt"])
    # Ensure East Asia font is set to the same font.
    style.element.rPr.rFonts.set(qn("w:eastAsia"), DEFAULTS["font_name"])

    p_format = style.paragraph_format
    p_format.line_spacing = DEFAULTS["line_spacing"]
    p_format.first_line_indent = Cm(DEFAULTS["first_line_indent_cm"])
    p_format.space_before = Pt(0)
    p_format.space_after = Pt(0)


def add_heading(doc: Document, level: int, text: str) -> None:
    p = doc.add_paragraph()
    run = p.add_run(text.strip())
    run.bold = True

    if level == 1:
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    else:
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT

    p.paragraph_format.first_line_indent = Cm(0)
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(6)


def add_paragraph(doc: Document, text: str) -> None:
    p = doc.add_paragraph(text.strip())
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY


def add_list_item(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    p.add_run("- " + text.strip())
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.first_line_indent = Cm(0)


def normalize_row(row: str) -> list[str]:
    cells = [c.strip() for c in row.strip().strip("|").split("|")]
    return cells


def add_table(doc: Document, rows: list[list[str]]) -> None:
    if not rows:
        return
    cols = max(len(r) for r in rows)
    table = doc.add_table(rows=len(rows), cols=cols)
    table.style = "Table Grid"

    for i, row in enumerate(rows):
        for j, cell_text in enumerate(row):
            cell = table.cell(i, j)
            cell.text = cell_text.strip()
            for p in cell.paragraphs:
                if i == 0:
                    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    for run in p.runs:
                        run.bold = True
                else:
                    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
                p.paragraph_format.first_line_indent = Cm(0)


def parse_markdown(lines: list[str], doc: Document) -> None:
    i = 0
    buffer: list[str] = []

    def flush_paragraph() -> None:
        nonlocal buffer
        if buffer:
            add_paragraph(doc, " ".join(buffer))
            buffer = []

    while i < len(lines):
        line = lines[i]

        if not line.strip():
            flush_paragraph()
            i += 1
            continue

        m_heading = HEADING_RE.match(line)
        if m_heading:
            flush_paragraph()
            level = len(m_heading.group(1))
            text = m_heading.group(2)
            add_heading(doc, level, text)
            i += 1
            continue

        m_list = LIST_RE.match(line)
        if m_list:
            flush_paragraph()
            add_list_item(doc, m_list.group(1))
            i += 1
            continue

        # Table detection: header row + separator row.
        if "|" in line and i + 1 < len(lines) and TABLE_SEPARATOR_RE.match(lines[i + 1]):
            flush_paragraph()
            rows: list[list[str]] = []
            rows.append(normalize_row(line))
            i += 2
            while i < len(lines) and "|" in lines[i] and lines[i].strip():
                rows.append(normalize_row(lines[i]))
                i += 1
            add_table(doc, rows)
            continue

        buffer.append(line.strip())
        i += 1

    flush_paragraph()


def main(argv: list[str]) -> int:
    if len(argv) < 2:
        print("Usage: md_to_docx_gost.py <input.md> [output.docx]")
        return 2

    input_path = Path(argv[1])
    output_path = Path(argv[2]) if len(argv) > 2 else input_path.with_suffix(".docx")

    if not input_path.exists():
        print(f"Input not found: {input_path}")
        return 2

    lines = input_path.read_text(encoding="utf-8").splitlines()
    doc = Document()
    set_document_style(doc)
    parse_markdown(lines, doc)
    doc.save(output_path)
    print(f"Saved: {output_path}")
    return 0


if __name__ == "__main__":
    raise SystemExit(main(sys.argv))
